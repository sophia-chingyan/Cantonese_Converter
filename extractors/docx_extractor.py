import io
from typing import Optional

from .base import Extractor, ExtractedDocument, ExtractionError


class DocxExtractor(Extractor):
    ext = "docx"

    def extract(self, raw_bytes: bytes, filename: Optional[str] = None) -> ExtractedDocument:
        try:
            import docx  # python-docx
        except ImportError as exc:
            raise ExtractionError("python-docx is not installed on the server.") from exc

        try:
            document = docx.Document(io.BytesIO(raw_bytes))
        except Exception as exc:
            raise ExtractionError(
                "Could not read this .docx file. It may be corrupted or not a real Word document."
            ) from exc

        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        if not paragraphs:
            raise ExtractionError("No text found in this document.")

        return ExtractedDocument(
            kind="plain",
            paragraphs=paragraphs,
            source_filename=filename,
            # D4: .docx input still produces .txt output.
            source_ext="docx",
        )
